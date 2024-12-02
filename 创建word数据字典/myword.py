#首先运行pip install python-docx来安装必要的库
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 创建一个新的Word文档
doc = Document()

# 添加标题
title = doc.add_heading('数据字典', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 定义表头
headers = ['字段名', '数据类型', '格式/约束', '描述', '范例值']

# 定义数据字典内容
data_dict = {
    '采购管理': {
        'suppliers': [
            ('supplier_id', 'INT', '主键, 自增', '供应商唯一标识符', '1, 2, 3, ...'),
            ('name', 'VARCHAR(255)', 'NOT NULL', '供应商名称', '"abc公司", "xyz供应商"'),
            ('contact_info', 'VARCHAR(255)', '', '联系信息（电话、邮箱等）', '"123-456-7890", "info@abc.com"'),
            ('address', 'TEXT', '', '供应商地址', '"北京市朝阳区XX路1号"'),
            ('bank_details', 'TEXT', '', '银行详情（用于支付）', '"银行名称: 工商银行, 账号: 1234567890"'),
            ('credit_rating', 'ENUM("a", "b", "c", "d")', 'NOT NULL', '信用评级', '"a", "b", "c", "d"'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'purchase_orders': [
            ('purchase_order_id', 'INT', '主键, 自增', '采购订单唯一标识符', '1, 2, 3, ...'),
            ('supplier_id', 'INT', '外键, 关联到suppliers', '供应商标识符', '1, 2, 3, ...'),
            ('order_date', 'DATE', 'NOT NULL', '订单创建日期', '"2024-01-01"'),
            ('delivery_date', 'DATE', '', '预计交货日期', '"2024-01-15"'),
            ('status', 'ENUM("已创建", "已确认", "已发货", "已接收", "已完成")', 'NOT NULL', '订单状态', '"已创建", "已确认"'),
            ('total_amount', 'DECIMAL(10, 2)', 'NOT NULL', '订单总金额', '1234.56, 7890.12'),
            ('created_by', 'INT', '', '创建人标识符', '1, 2, 3, ...'),
            ('approved_by', 'INT', '', '审批人标识符', '1, 2, 3, ...'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'purchase_order_items': [
            ('purchase_order_item_id', 'INT', '主键, 自增', '采购订单项唯一标识符', '1, 2, 3, ...'),
            ('purchase_order_id', 'INT', '外键, 关联到purchase_orders', '采购订单标识符', '1, 2, 3, ...'),
            ('product_id', 'INT', '外键, 关联到products', '产品标识符', '1, 2, 3, ...'),
            ('quantity', 'INT', 'NOT NULL', '数量', '10, 50, 100'),
            ('unit_price', 'DECIMAL(10, 2)', 'NOT NULL', '单价', '12.34, 56.78'),
            ('discount', 'DECIMAL(5, 2)', 'DEFAULT 0.00', '折扣', '0.00, 5.00, 10.00'),
            ('tax', 'DECIMAL(5, 2)', 'DEFAULT 0.00', '税额', '0.00, 5.00, 10.00'),
            ('sub_total', 'DECIMAL(10, 2)', 'NOT NULL', '小计', '123.45, 678.90'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ]
    },
    '库存管理': {
        'products': [
            ('product_id', 'INT', '主键, 自增', '产品唯一标识符', '1, 2, 3, ...'),
            ('sku', 'VARCHAR(50)', 'UNIQUE, NOT NULL', '库存保有单位', '"sku123", "sku456"'),
            ('name', 'VARCHAR(255)', 'NOT NULL', '产品名称', '"产品a", "产品b"'),
            ('description', 'TEXT', '', '产品描述', '"这是一个很好的产品"'),
            ('category', 'VARCHAR(100)', 'NOT NULL', '产品类别', '"电子产品", "办公用品"'),
            ('unit_of_measure', 'VARCHAR(50)', 'NOT NULL', '计量单位', '"件", "箱", "公斤"'),
            ('reorder_level', 'INT', 'DEFAULT 0', '再订购水平', '10, 50, 100'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'inventory_transactions': [
            ('transaction_id', 'INT', '主键, 自增', '交易唯一标识符', '1, 2, 3, ...'),
            ('product_id', 'INT', '外键, 关联到products', '产品标识符', '1, 2, 3, ...'),
            ('transaction_type', 'ENUM("入库", "出库", "调整")', 'NOT NULL', '交易类型', '"入库", "出库", "调整"'),
            ('quantity', 'INT', 'NOT NULL', '数量', '10, 50, 100'),
            ('price', 'DECIMAL(10, 2)', 'NOT NULL', '单价', '12.34, 56.78'),
            ('transaction_date', 'DATE', 'NOT NULL', '交易日期', '"2024-01-01"'),
            ('reference', 'VARCHAR(255)', '', '关联参考（如purchase_order_id或sales_order_id）', '"PO123", "SO456"'),
            ('created_by', 'INT', '', '创建人标识符', '1, 2, 3, ...'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'inventory_snapshots': [
            ('snapshot_id', 'INT', '主键, 自增', '快照唯一标识符', '1, 2, 3, ...'),
            ('product_id', 'INT', '外键, 关联到products', '产品标识符', '1, 2, 3, ...'),
            ('available_quantity', 'INT', 'NOT NULL', '可用数量', '10, 50, 100'),
            ('snapshot_date', 'DATE', 'NOT NULL', '快照日期', '"2024-01-01"'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"')
        ]
    },
    '销售管理': {
        'customers': [
            ('customer_id', 'INT', '主键, 自增', '客户唯一标识符', '1, 2, 3, ...'),
            ('name', 'VARCHAR(255)', 'NOT NULL', '客户名称', '"张三", "李四"'),
            ('contact_info', 'VARCHAR(255)', '', '联系信息（电话、邮箱等）', '"123-456-7890", "info@customer.com"'),
            ('address', 'TEXT', '', '客户地址', '"上海市浦东新区XX路1号"'),
            ('preferred_payment_method', 'VARCHAR(100)', 'NOT NULL', '偏好的支付方式', '"信用卡", "支付宝"'),
            ('credit_limit', 'DECIMAL(10, 2)', 'DEFAULT 0.00', '信用限额', '1000.00, 5000.00'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'sales_orders': [
            ('sales_order_id', 'INT', '主键, 自增', '销售订单唯一标识符', '1, 2, 3, ...'),
            ('customer_id', 'INT', '外键, 关联到customers', '客户标识符', '1, 2, 3, ...'),
            ('order_date', 'DATE', 'NOT NULL', '订单创建日期', '"2024-01-01"'),
            ('required_date', 'DATE', '', '需求日期', '"2024-01-15"'),
            ('shipped_date', 'DATE', '', '发货日期', '"2024-01-10"'),
            ('status', 'ENUM("处理中", "已发货", "已完成", "已取消")', 'NOT NULL', '订单状态', '"处理中", "已发货"'),
            ('total_amount', 'DECIMAL(10, 2)', 'NOT NULL', '订单总金额', '1234.56, 7890.12'),
            ('created_by', 'INT', '', '创建人标识符', '1, 2, 3, ...'),
            ('approved_by', 'INT', '', '审批人标识符', '1, 2, 3, ...'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'sales_order_items': [
            ('sales_order_item_id', 'INT', '主键, 自增', '销售订单项唯一标识符', '1, 2, 3, ...'),
            ('sales_order_id', 'INT', '外键, 关联到sales_orders', '销售订单标识符', '1, 2, 3, ...'),
            ('product_id', 'INT', '外键, 关联到products', '产品标识符', '1, 2, 3, ...'),
            ('quantity', 'INT', 'NOT NULL', '数量', '10, 50, 100'),
            ('unit_price', 'DECIMAL(10, 2)', 'NOT NULL', '单价', '12.34, 56.78'),
            ('discount', 'DECIMAL(5, 2)', 'DEFAULT 0.00', '折扣', '0.00, 5.00, 10.00'),
            ('tax', 'DECIMAL(5, 2)', 'DEFAULT 0.00', '税额', '0.00, 5.00, 10.00'),
            ('sub_total', 'DECIMAL(10, 2)', 'NOT NULL', '小计', '123.45, 678.90'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ],
        'shipments': [
            ('shipment_id', 'INT', '主键, 自增', '发货唯一标识符', '1, 2, 3, ...'),
            ('sales_order_id', 'INT', '外键, 关联到sales_orders', '销售订单标识符', '1, 2, 3, ...'),
            ('shipment_date', 'DATE', 'NOT NULL', '发货日期', '"2024-01-10"'),
            ('carrier', 'VARCHAR(255)', 'NOT NULL', '物流公司', '"顺丰快递", "圆通速递"'),
            ('tracking_number', 'VARCHAR(255)', 'UNIQUE, NOT NULL', '追踪号码', '"SF1234567890", "YTO9876543210"'),
            ('shipping_cost', 'DECIMAL(10, 2)', 'NOT NULL', '运费', '12.34, 56.78'),
            ('created_date', 'DATETIME', 'DEFAULT CURRENT_TIMESTAMP', '创建日期', '"2024-01-01 10:00:00"'),
            ('updated_date', 'DATETIME', 'ON UPDATE CURRENT_TIMESTAMP', '最后更新日期', '"2024-01-02 11:00:00"')
        ]
    }
}

# 遍历数据字典并生成表格
for section, tables in data_dict.items():
    # 添加章节标题
    section_heading = doc.add_heading(section, level=2)
    section_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for table_name, fields in tables.items():
        # 添加表标题
        table_heading = doc.add_heading(table_name, level=3)
        table_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 创建表格
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'  # 设置表格样式
        table.autofit = False
        table.allow_autofit = False

        # 设置表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
            hdr_cells[i].width = Inches(2.5)  # 设置列宽

        # 填充表格内容
        for field in fields:
            row_cells = table.add_row().cells
            for i, value in enumerate(field):
                row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
                row_cells[i].width = Inches(2.5)  # 设置列宽

        # 设置表格样式
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_TABLE_ALIGNMENT.CENTER  # 垂直居中对齐
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 水平左对齐

        # 添加段落间距
        doc.add_paragraph('')

# 保存文档
doc.save('data_dictionary.docx')

print("数据字典已生成并保存为 data_dictionary.docx")